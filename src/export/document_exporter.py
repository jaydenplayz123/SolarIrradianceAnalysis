"""
Document exporter for Solar Irradiance Analysis
Sistema de exportación de documentos para Análisis de Irradiación Solar
"""

from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np

# DOCX export
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors

# PDF export
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from src.analysis.analysis_functions import (
    format_location_for_display,
    is_combined_dataset,
)
from src.data.data_sets import (
    get_all_datasets_period,
)
from src.utils.i18n import i18n


class DocumentExporter:
    """
    Class for exporting analysis results to DOCX and PDF documents
    """

    def __init__(self, output_dir="reports"):
        """
        Initialize the document exporter

        Args:
            output_dir (str): Directory where reports will be saved
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Create subdirectories for images
        self.images_dir = self.output_dir / "images"
        self.images_dir.mkdir(exist_ok=True)

        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    def save_plot_as_image(self, figure, filename):
        """
        Save a matplotlib figure as an image file

        Args:
            figure: matplotlib figure object
            filename (str): name for the image file (without extension)

        Returns:
            Path: path to the saved image
        """
        image_path = self.images_dir / f"{filename}_{self.timestamp}.png"
        figure.savefig(image_path, dpi=300, bbox_inches="tight", facecolor="white")
        plt.close(figure)  # Close figure to free memory
        return image_path

    def create_plots_for_dataset(
        self, results, dataset_name, months, location="Location data"
    ):
        """
        Create and save plots for a specific dataset

        Args:
            results (dict): Analysis results for the dataset
            dataset_name (str): Name of the dataset
            months (list): List of month names
            location (str): Location information for the dataset

        Returns:
            dict: Dictionary with paths to saved plot images
        """
        print(f"   📊 {i18n.t_print('creating_plots')} {dataset_name}...")

        plot_paths = {}

        try:
            # Create main analysis plot
            fig, axes = plt.subplots(2, 2, figsize=(16, 12))

            # Use figtext for better control over title positioning and line breaks
            title_line1 = i18n.t_plot("monthly_time_series")
            title_line2 = f"{dataset_name} - ALLSKY_SFC_SW_DWN"

            # Add title with proper spacing - positioned at top center
            fig.text(
                0.5,
                0.96,
                title_line1,
                ha="center",
                va="top",
                fontsize=16,
                fontweight="bold",
            )
            fig.text(
                0.5,
                0.93,
                title_line2,
                ha="center",
                va="top",
                fontsize=14,
                fontweight="bold",
            )

            # Add location information - special handling for combined datasets
            if is_combined_dataset(dataset_name, location):
                formatted_location = format_location_for_display(location, True, i18n)
                fig.text(
                    0.5,
                    0.90,
                    formatted_location,
                    ha="center",
                    va="top",
                    fontsize=10,
                    fontweight="normal",
                    style="italic",
                    color="gray",
                )
            else:
                # Single location
                formatted_location = format_location_for_display(location, False, i18n)
                fig.text(
                    0.5,
                    0.90,
                    formatted_location,
                    ha="center",
                    va="top",
                    fontsize=10,
                    fontweight="normal",
                    style="italic",
                    color="gray",
                )

            # Get plot months (for plot language)
            plot_months = i18n.get_months_plot()

            # Extract data from results
            df = results["dataframe"]
            annual_mean = results["annual_mean"]
            climatology = results["climatology"]
            anomalies = results["anomalies"]
            annual_anomalies = results["annual_anomalies"]
            z = results["trend_coefficients"]
            p = results["trend_function"]

            # 1. MONTHLY TIME SERIES
            ax1 = axes[0, 0]
            colors = plt.cm.tab20(np.linspace(0, 1, len(months)))
            for i, month in enumerate(months):
                ax1.plot(
                    df.index,
                    df[month],
                    label=plot_months[i],
                    linewidth=1,
                    alpha=0.7,
                    color=colors[i],
                )

            ax1.set_title(
                i18n.t_plot("monthly_time_series_subplot"),
                fontweight="bold",
                loc="left",
            )
            ax1.set_xlabel(i18n.t_plot("year"))
            ax1.set_ylabel(i18n.t_plot("irradiance_kwh"))
            ax1.grid(True, alpha=0.3)
            ax1.legend(
                ncol=4, fontsize=8, loc="upper center", bbox_to_anchor=(0.5, -0.15)
            )
            ax1.set_xlim(df.index.min(), df.index.max())

            # Dynamic Y-axis limits based on data range
            y_min = df.values.min()
            y_max = df.values.max()
            y_margin = (y_max - y_min) * 0.1  # 10% margin
            ax1.set_ylim(y_min - y_margin, y_max + y_margin)

            # 2. ANNUAL MEAN
            ax2 = axes[0, 1]
            ax2.plot(
                annual_mean.index,
                annual_mean.values,
                linewidth=2.5,
                color="red",
                marker="o",
                markersize=3,
            )
            ax2.set_title(
                i18n.t_plot("annual_mean_subplot"), fontweight="bold", loc="left"
            )
            ax2.set_xlabel(i18n.t_plot("year"))
            ax2.set_ylabel(i18n.t_plot("mean_annual_irradiance"))
            ax2.grid(True, alpha=0.3)
            ax2.set_xlim(annual_mean.index.min(), annual_mean.index.max())

            # Linear trend
            x_trend = annual_mean.index
            y_trend = p(x_trend)
            ax2.plot(
                x_trend,
                y_trend,
                "r--",
                alpha=0.8,
                linewidth=1.5,
                label=f'{i18n.t_plot("trend_label")}: {z[0]:.4f} {i18n.t_plot("kwh_per_m2_year")}',
            )
            ax2.legend()

            # 3. MONTHLY CLIMATOLOGY
            ax3 = axes[1, 0]
            # Plot climatology
            ax3.plot(
                plot_months,
                climatology.values,
                marker="o",
                linewidth=2.5,
                color="green",
                markerfacecolor="darkgreen",
                markersize=6,
            )

            # Get dynamic period for subtitle
            from src.data.data_sets import get_dataset_years

            data_dict_for_years = {
                year: results["dataframe"].loc[year].values
                for year in results["dataframe"].index
            }
            start_year, end_year = get_dataset_years(data_dict_for_years)
            climatology_title = (
                f"{i18n.t_plot('climatology_subplot')} - {start_year}-{end_year}"
            )
            ax3.set_title(climatology_title, fontweight="bold", loc="left")
            ax3.set_xlabel(i18n.t_plot("month"))
            ax3.set_ylabel(i18n.t_plot("mean_irradiance"))
            ax3.grid(True, alpha=0.3)

            # Dynamic Y-axis limits for climatology
            clim_min = climatology.min()
            clim_max = climatology.max()
            clim_margin = (
                clim_max - clim_min
            ) * 0.15  # 15% margin for better visibility
            ax3.set_ylim(clim_min - clim_margin, clim_max + clim_margin)
            ax3.fill_between(plot_months, climatology.values, alpha=0.2, color="green")

            # 4. MONTHLY ANOMALIES
            ax4 = axes[1, 1]

            # Monthly anomalies
            for i, month in enumerate(months):
                ax4.plot(
                    anomalies.index,
                    anomalies[month],
                    alpha=0.2,
                    linewidth=0.5,
                    color=colors[i],
                )

            # Annual mean anomalies
            ax4.plot(
                annual_anomalies.index,
                annual_anomalies.values,
                linewidth=2.5,
                color="black",
                label=i18n.t_plot("annual_mean"),
                marker="o",
                markersize=3,
            )

            # 5-year rolling mean
            rolling_mean = annual_anomalies.rolling(window=5, center=True).mean()
            ax4.plot(
                rolling_mean.index,
                rolling_mean.values,
                linewidth=2,
                color="blue",
                label=i18n.t_plot("rolling_mean_5_years"),
            )

            ax4.set_title(
                i18n.t_plot("anomalies_subplot"), fontweight="bold", loc="left"
            )
            ax4.set_xlabel(i18n.t_plot("year"))
            ax4.set_ylabel(i18n.t_plot("anomaly_kwh"))
            ax4.grid(True, alpha=0.3)
            ax4.axhline(y=0, color="black", linestyle="-", alpha=0.7)

            # Set x-limits
            ax4.set_xlim(anomalies.index.min(), anomalies.index.max())

            # Dynamic Y-axis limits for anomalies
            anom_min = annual_anomalies.min()
            anom_max = annual_anomalies.max()

            # Include monthly anomalies in range calculation
            monthly_min = anomalies.values.min()
            monthly_max = anomalies.values.max()
            anom_min = min(anom_min, monthly_min)
            anom_max = max(anom_max, monthly_max)

            anom_margin = (
                max(abs(anom_min), abs(anom_max)) * 0.2
            )  # 20% margin, symmetric
            ax4.set_ylim(-anom_margin - abs(anom_min), anom_margin + abs(anom_max))
            ax4.legend()

            plt.tight_layout()
            # Adjust spacing to accommodate the title and location - more space for combined datasets
            if is_combined_dataset(dataset_name, location):
                plt.subplots_adjust(top=0.85)  # More space for multi-line location
            else:
                plt.subplots_adjust(top=0.87)  # Standard space for single location

            # Save main plot
            main_plot_path = self.save_plot_as_image(
                fig, f"main_analysis_{dataset_name.lower().replace(' ', '_')}"
            )
            plot_paths["main_plot"] = main_plot_path

            # Create decade analysis plot
            fig2 = plt.figure(figsize=(12, 6))
            decades = [
                (start_year, start_year + 9),
                (start_year + 10, start_year + 19),
                (start_year + 20, start_year + 29),
                (start_year + 30, end_year),
            ]
            decade_colors = ["blue", "green", "orange", "red"]

            for i, (dec_start, dec_end) in enumerate(decades):
                # Filter decade data
                dec_data = annual_mean[
                    (annual_mean.index >= dec_start) & (annual_mean.index <= dec_end)
                ]
                if len(dec_data) > 0:
                    plt.plot(
                        dec_data.index,
                        dec_data.values,
                        "o-",
                        color=decade_colors[i],
                        label=f"{dec_start}-{dec_end}",
                        alpha=0.7,
                        markersize=4,
                    )

            # Add trend line to decades plot
            plt.plot(
                annual_mean.index,
                p(annual_mean.index),
                "k--",
                linewidth=2,
                label=f'{i18n.t_plot("general_trend")} ({z[0]:.4f}/{i18n.t_plot("year")})',
            )

            # Title with bold formatting - consistent with main plot format
            decade_title_line1 = i18n.t_plot("annual_mean_irradiance_evolution")
            decade_title_line2 = f"{dataset_name} ({start_year}-{end_year})"

            # Add title with better spacing for decade plot
            fig2.text(
                0.5,
                0.96,
                decade_title_line1,
                ha="center",
                va="top",
                fontsize=16,
                fontweight="bold",
            )
            fig2.text(
                0.5,
                0.91,
                decade_title_line2,
                ha="center",
                va="top",
                fontsize=14,
                fontweight="bold",
            )

            # Add location information with proper spacing
            if is_combined_dataset(dataset_name, location):
                formatted_location = format_location_for_display(location, True, i18n)
                fig2.text(
                    0.5,
                    0.87,
                    formatted_location,
                    ha="center",
                    va="top",
                    fontsize=10,
                    fontweight="normal",
                    style="italic",
                    color="gray",
                )
            else:
                # Single location
                formatted_location = format_location_for_display(location, False, i18n)
                fig2.text(
                    0.5,
                    0.87,
                    formatted_location,
                    ha="center",
                    va="top",
                    fontsize=10,
                    fontweight="normal",
                    style="italic",
                    color="gray",
                )
            plt.xlabel(i18n.t_plot("year"))
            plt.ylabel(i18n.t_plot("annual_mean_irradiance"))
            plt.grid(True, alpha=0.3)
            plt.legend()
            plt.tight_layout()
            # Adjust spacing to accommodate the title and location - more space for decade plot
            if is_combined_dataset(dataset_name, location):
                plt.subplots_adjust(top=0.82)  # More space for multi-line location
            else:
                plt.subplots_adjust(top=0.84)  # Standard space for single location

            # Save decade plot
            decade_plot_path = self.save_plot_as_image(
                fig2, f"decade_analysis_{dataset_name.lower().replace(' ', '_')}"
            )
            plot_paths["decade_plot"] = decade_plot_path

        except Exception as e:
            print(f"   ⚠️ Error creating plots for {dataset_name}: {e}")

        return plot_paths

    def create_comparative_plot(self, all_results):
        """
        Create comparative summary plot for all datasets

        Args:
            all_results (dict): Dictionary with analysis results for all datasets

        Returns:
            Path: path to the saved comparative plot image
        """
        print(f"   📊 {i18n.t_print('creating_plots')} - Comparative Summary...")

        try:
            import matplotlib.pyplot as plt
            import numpy as np

            # Prepare data for comparison
            dataset_names = list(all_results.keys())
            mean_irradiances = []
            trends = []
            variabilities = []

            for name, results in all_results.items():
                annual_mean = results["annual_mean"]

                # Calculate statistics from pandas Series
                mean_val = annual_mean.mean()
                std_val = annual_mean.std()

                mean_irradiances.append(mean_val)
                trends.append(
                    results["trend_coefficients"][0] * 1000
                )  # Convert to per 1000 years for better visualization
                cv = (std_val / mean_val) * 100
                variabilities.append(cv)

            # Create comparative plot
            fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(16, 12))
            fig.suptitle(
                f'{i18n.t_plot("comparative_summary")} - {i18n.t_plot("multiple_datasets")}',
                fontsize=16,
                fontweight="bold",
            )

            colors = plt.cm.Set3(np.linspace(0, 1, len(dataset_names)))

            # 1. Mean Irradiance Comparison
            bars1 = ax1.bar(
                dataset_names,
                mean_irradiances,
                color=colors,
                alpha=0.7,
                edgecolor="black",
            )
            ax1.set_title(
                f'{i18n.t_plot("mean_irradiance")} {i18n.t_plot("comparison")}',
                fontweight="bold",
            )
            ax1.set_ylabel(
                f'{i18n.t_plot("irradiance")} ({i18n.t_plot("kwh_per_m2_day")})'
            )
            ax1.grid(True, alpha=0.3)
            ax1.tick_params(axis="x", rotation=45)

            # Add value labels on bars
            for bar, value in zip(bars1, mean_irradiances):
                height = bar.get_height()
                ax1.text(
                    bar.get_x() + bar.get_width() / 2.0,
                    height + 0.01,
                    f"{value:.3f}",
                    ha="center",
                    va="bottom",
                    fontsize=10,
                )

            # 2. Trend Comparison
            bars2 = ax2.bar(
                dataset_names, trends, color=colors, alpha=0.7, edgecolor="black"
            )
            ax2.set_title(
                f'{i18n.t_plot("trend")} {i18n.t_plot("comparison")}', fontweight="bold"
            )
            ax2.set_ylabel(
                f'{i18n.t_plot("trend")} ({i18n.t_plot("kwh_per_m2_day")}/1000 {i18n.t_plot("years")})'
            )
            ax2.grid(True, alpha=0.3)
            ax2.tick_params(axis="x", rotation=45)
            ax2.axhline(y=0, color="red", linestyle="--", alpha=0.7)

            # Add value labels on bars
            for bar, value in zip(bars2, trends):
                height = bar.get_height()
                y_pos = height + 0.1 if height >= 0 else height - 0.1
                ax2.text(
                    bar.get_x() + bar.get_width() / 2.0,
                    y_pos,
                    f"{value:.2f}",
                    ha="center",
                    va="bottom" if height >= 0 else "top",
                    fontsize=10,
                )

            # 3. Variability Comparison
            bars3 = ax3.bar(
                dataset_names, variabilities, color=colors, alpha=0.7, edgecolor="black"
            )
            ax3.set_title(
                f'{i18n.t_plot("variability")} {i18n.t_plot("comparison")}',
                fontweight="bold",
            )
            ax3.set_ylabel(f'{i18n.t_plot("coefficient_variation")} (%)')
            ax3.grid(True, alpha=0.3)
            ax3.tick_params(axis="x", rotation=45)

            # Add value labels on bars
            for bar, value in zip(bars3, variabilities):
                height = bar.get_height()
                ax3.text(
                    bar.get_x() + bar.get_width() / 2.0,
                    height + 0.02,
                    f"{value:.2f}%",
                    ha="center",
                    va="bottom",
                    fontsize=10,
                )

            # 4. Combined scatter plot (Mean vs Variability, colored by trend)
            scatter = ax4.scatter(
                mean_irradiances,
                variabilities,
                c=trends,
                s=200,
                alpha=0.7,
                cmap="RdYlBu_r",
                edgecolors="black",
                linewidth=2,
            )
            ax4.set_title(
                f'{i18n.t_plot("mean_irradiance")} vs {i18n.t_plot("variability")}',
                fontweight="bold",
            )
            ax4.set_xlabel(
                f'{i18n.t_plot("mean_irradiance")} ({i18n.t_plot("kwh_per_m2_day")})'
            )
            ax4.set_ylabel(f'{i18n.t_plot("coefficient_variation")} (%)')
            ax4.grid(True, alpha=0.3)

            # Add colorbar for trend
            cbar = plt.colorbar(scatter, ax=ax4)
            cbar.set_label(
                f'{i18n.t_plot("trend")} ({i18n.t_plot("kwh_per_m2_day")}/1000 {i18n.t_plot("years")})'
            )

            # Add dataset labels to scatter points
            for i, name in enumerate(dataset_names):
                ax4.annotate(
                    name,
                    (mean_irradiances[i], variabilities[i]),
                    xytext=(5, 5),
                    textcoords="offset points",
                    fontsize=9,
                )

            plt.tight_layout()
            plt.subplots_adjust(top=0.93)

            # Save comparative plot
            comparative_plot_path = self.save_plot_as_image(fig, "comparative_summary")
            return comparative_plot_path

        except Exception as e:
            print(f"   ⚠️ Error creating comparative plot: {e}")
            return None

    def add_complete_analysis_to_docx(self, doc, results, dataset_name):
        """
        Add complete detailed analysis to DOCX document (matching console output)

        Args:
            doc: Document object
            results (dict): Analysis results
            dataset_name (str): Name of the dataset
        """
        # Data is already in pandas format - no conversion needed
        df = results["dataframe"]
        annual_mean = results["annual_mean"]
        climatology = results["climatology"]
        anomalies = results["anomalies"]
        annual_anomalies = results["annual_anomalies"]
        trend_coefficients = results["trend_coefficients"]

        # Get months for display
        months = i18n.get_months_print()

        # Basic dataset info
        doc.add_heading(f"{i18n.t_print('analysis').upper()} - {dataset_name}", level=2)

        basic_info = doc.add_paragraph()
        basic_info.add_run(
            f"{i18n.t_print('period')}: {df.index.min()}-{df.index.max()}\n"
        )
        basic_info.add_run(
            f"Dimension: {len(df)} {i18n.t_print('years')} × 12 {i18n.t_print('months').lower()}\n"
        )
        basic_info.add_run(
            f"Global range: {df.values.min():.2f} - {df.values.max():.2f} {i18n.t_print('kwh_per_m2_day')}"
        )

        doc.add_paragraph()  # Spacer

        # 1. GLOBAL STATISTICS
        doc.add_heading(f"1. {i18n.t_print('global_statistics')}", level=3)

        global_stats = doc.add_paragraph()
        global_stats.add_run(
            f"   • {i18n.t_print('period')}: {df.index.min()}-{df.index.max()} ({len(df)} {i18n.t_print('years')})\n"
        )
        global_stats.add_run(
            f"   • {i18n.t_print('global_mean')}: {df.values.mean():.4f} {i18n.t_print('kwh_per_m2_day')}\n"
        )
        global_stats.add_run(
            f"   • {i18n.t_print('standard_deviation')}: {df.values.std():.4f} {i18n.t_print('kwh_per_m2_day')}\n"
        )
        global_stats.add_run(
            f"   • {i18n.t_print('total_range')}: {df.values.min():.4f} - {df.values.max():.4f} {i18n.t_print('kwh_per_m2_day')}"
        )

        # 2. ANNUAL STATISTICS
        doc.add_heading(f"2. {i18n.t_print('annual_statistics')}", level=3)

        # Get statistics
        annual_stats_data = {
            "mean": annual_mean.mean(),
            "std": annual_mean.std(),
            "min": annual_mean.min(),
            "max": annual_mean.max(),
            "idxmax": annual_mean.idxmax(),
            "idxmin": annual_mean.idxmin(),
        }

        annual_stats = doc.add_paragraph()
        annual_stats.add_run(
            f"   • {i18n.t_print('average_annual_mean')}: {annual_stats_data['mean']:.4f} ± {annual_stats_data['std']:.4f} {i18n.t_print('kwh_per_m2_day')}\n"
        )
        annual_stats.add_run(
            f"   • {i18n.t_print('year_highest_irradiance')}: {annual_stats_data['idxmax']} ({annual_stats_data['max']:.4f})\n"
        )
        annual_stats.add_run(
            f"   • {i18n.t_print('year_lowest_irradiance')}: {annual_stats_data['idxmin']} ({annual_stats_data['min']:.4f})\n"
        )
        annual_stats.add_run(
            f"   • {i18n.t_print('linear_trend')}: {trend_coefficients[0]:.6f} {i18n.t_print('kwh_per_m2_day')}/{i18n.t_print('year')}"
        )

        # 3. MONTHLY CLIMATOLOGY with table
        from src.data.data_sets import get_dataset_years

        # Convert df back to proper format for get_dataset_years
        data_dict_for_years = {
            year: results["dataframe"].loc[year].values
            for year in results["dataframe"].index
        }
        start_year, end_year = get_dataset_years(data_dict_for_years)
        doc.add_heading(
            f"3. {i18n.t_print('monthly_climatology')} ({start_year}-{end_year})",
            level=3,
        )

        # Create climatology table
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = i18n.t_print("month_header")
        header_cells[1].text = i18n.t_print("average_header")
        header_cells[2].text = i18n.t_print("maximum_header")
        header_cells[3].text = i18n.t_print("minimum_header")
        header_cells[4].text = i18n.t_print("std_header")

        # Data rows
        for i, month in enumerate(months):
            month_data = df[month]
            row_cells = table.add_row().cells
            row_cells[0].text = month
            # Access climatology data
            clim_value = climatology.iloc[i]
            row_cells[1].text = f"{clim_value:.4f}"

            # Calculate month statistics
            row_cells[2].text = f"{month_data.max():.4f}"
            row_cells[3].text = f"{month_data.min():.4f}"
            row_cells[4].text = f"{month_data.std():.4f}"

        # 4. ANOMALY ANALYSIS
        doc.add_heading(f"4. {i18n.t_print('anomaly_analysis')}", level=3)

        # Get anomaly statistics
        anomaly_stats_data = {
            "mean": annual_anomalies.mean(),
            "std": annual_anomalies.std(),
            "min": annual_anomalies.min(),
            "max": annual_anomalies.max(),
            "idxmax": annual_anomalies.idxmax(),
            "idxmin": annual_anomalies.idxmin(),
        }

        # Count positive/negative anomalies
        positive_count = sum(annual_anomalies > 0)
        negative_count = sum(annual_anomalies < 0)

        anomaly_stats = doc.add_paragraph()
        anomaly_stats.add_run(
            f"   • {i18n.t_print('year_highest_positive_anomaly')}: {anomaly_stats_data['idxmax']} ({anomaly_stats_data['max']:.4f})\n"
        )
        anomaly_stats.add_run(
            f"   • {i18n.t_print('year_highest_negative_anomaly')}: {anomaly_stats_data['idxmin']} ({anomaly_stats_data['min']:.4f})\n"
        )
        anomaly_stats.add_run(
            f"   • {i18n.t_print('years_positive_anomaly')}: {positive_count}\n"
        )
        anomaly_stats.add_run(
            f"   • {i18n.t_print('years_negative_anomaly')}: {negative_count}\n"
        )
        anomaly_stats.add_run(
            f"   • {i18n.t_print('standard_deviation_anomalies')}: {anomaly_stats_data['std']:.4f}"
        )

        # 5. SEASONALITY
        doc.add_heading(f"5. {i18n.t_print('seasonality')}", level=3)

        import numpy as np

        # Process climatology
        clim_stats = {
            "mean": climatology.mean(),
            "std": climatology.std(),
            "min": climatology.min(),
            "max": climatology.max(),
            "idxmax": climatology.idxmax(),
            "idxmin": climatology.idxmin(),
        }
        seasonal_amplitude = clim_stats["max"] - clim_stats["min"]

        # Find max and min months
        max_month = months[np.argmax(climatology.values)]
        min_month = months[np.argmin(climatology.values)]

        seasonality_stats = doc.add_paragraph()
        seasonality_stats.add_run(
            f"   • {i18n.t_print('month_highest_irradiance')}: {max_month} ({clim_stats['max']:.4f} {i18n.t_print('kwh_per_m2_day')})\n"
        )
        seasonality_stats.add_run(
            f"   • {i18n.t_print('month_lowest_irradiance')}: {min_month} ({clim_stats['min']:.4f} {i18n.t_print('kwh_per_m2_day')})\n"
        )
        seasonality_stats.add_run(
            f"   • {i18n.t_print('seasonal_amplitude')}: {seasonal_amplitude:.4f} {i18n.t_print('kwh_per_m2_day')}"
        )

        # 6. INTERANNUAL VARIABILITY
        doc.add_heading(f"6. {i18n.t_print('interannual_variability')}", level=3)

        cv = (annual_stats_data["std"] / annual_stats_data["mean"]) * 100
        interannual_range = annual_stats_data["max"] - annual_stats_data["min"]

        variability_stats = doc.add_paragraph()
        variability_stats.add_run(
            f"   • {i18n.t_print('annual_coefficient_variation')}: {cv:.2f}%\n"
        )
        variability_stats.add_run(
            f"   • {i18n.t_print('interannual_range')}: {interannual_range:.4f} {i18n.t_print('kwh_per_m2_day')}"
        )

        # 7. FINAL SUMMARY
        doc.add_heading(f"7. {i18n.t_print('final_summary')}", level=3)

        trend_direction = (
            i18n.t_print("negative")
            if trend_coefficients[0] < 0
            else i18n.t_print("positive")
        )

        summary_stats = doc.add_paragraph()
        summary_stats.add_run(
            f"   • {i18n.t_print('general_trend')}: {trend_direction.upper()}\n"
        )
        summary_stats.add_run(
            f"   • {i18n.t_print('seasonal_variability')}: {seasonal_amplitude:.4f} {i18n.t_print('kwh_per_m2_day')}\n"
        )
        summary_stats.add_run(
            f"   • {i18n.t_print('interannual_stability')}: {cv:.2f}% {i18n.t_print('variation')}"
        )

    def add_complete_analysis_to_pdf(
        self, story, styles, heading_style, results, dataset_name
    ):
        """
        Add complete detailed analysis to PDF document (matching console output)

        Args:
            story: PDF story list
            styles: PDF styles
            heading_style: PDF heading style
            results (dict): Analysis results
            dataset_name (str): Name of the dataset
        """
        # Data is already in pandas format - no conversion needed
        df = results["dataframe"]
        annual_mean = results["annual_mean"]
        climatology = results["climatology"]
        anomalies = results["anomalies"]
        annual_anomalies = results["annual_anomalies"]
        trend_coefficients = results["trend_coefficients"]

        # Get months for display
        months = i18n.get_months_print()

        # Basic dataset info
        story.append(
            Paragraph(
                f"{i18n.t_print('analysis').upper()} - {dataset_name}", heading_style
            )
        )

        basic_info_text = f"""
        {i18n.t_print('period')}: {df.index.min()}-{df.index.max()}<br/>
        Dimension: {len(df)} {i18n.t_print('years')} × 12 {i18n.t_print('months').lower()}<br/>
        Global range: {df.values.min():.2f} - {df.values.max():.2f} {i18n.t_print('kwh_per_m2_day')}
        """
        story.append(Paragraph(basic_info_text, styles["Normal"]))
        story.append(Spacer(1, 12))

        # 1. GLOBAL STATISTICS
        story.append(
            Paragraph(f"1. {i18n.t_print('global_statistics')}", heading_style)
        )

        global_stats_text = f"""
        • {i18n.t_print('period')}: {df.index.min()}-{df.index.max()} ({len(df)} {i18n.t_print('years')})<br/>
        • {i18n.t_print('global_mean')}: {df.values.mean():.4f} {i18n.t_print('kwh_per_m2_day')}<br/>
        • {i18n.t_print('standard_deviation')}: {df.values.std():.4f} {i18n.t_print('kwh_per_m2_day')}<br/>
        • {i18n.t_print('total_range')}: {df.values.min():.4f} - {df.values.max():.4f} {i18n.t_print('kwh_per_m2_day')}
        """
        story.append(Paragraph(global_stats_text, styles["Normal"]))
        story.append(Spacer(1, 12))

        # 2. ANNUAL STATISTICS
        story.append(
            Paragraph(f"2. {i18n.t_print('annual_statistics')}", heading_style)
        )

        # Get statistics for PDF
        annual_stats_data = {
            "mean": annual_mean.mean(),
            "std": annual_mean.std(),
            "min": annual_mean.min(),
            "max": annual_mean.max(),
            "idxmax": annual_mean.idxmax(),
            "idxmin": annual_mean.idxmin(),
        }

        annual_stats_text = f"""
        • {i18n.t_print('average_annual_mean')}: {annual_stats_data['mean']:.4f} ± {annual_stats_data['std']:.4f} {i18n.t_print('kwh_per_m2_day')}<br/>
        • {i18n.t_print('year_highest_irradiance')}: {annual_stats_data['idxmax']} ({annual_stats_data['max']:.4f})<br/>
        • {i18n.t_print('year_lowest_irradiance')}: {annual_stats_data['idxmin']} ({annual_stats_data['min']:.4f})<br/>
        • {i18n.t_print('linear_trend')}: {trend_coefficients[0]:.6f} {i18n.t_print('kwh_per_m2_day')}/{i18n.t_print('year')}
        """
        story.append(Paragraph(annual_stats_text, styles["Normal"]))
        story.append(Spacer(1, 12))

        # 3. MONTHLY CLIMATOLOGY with table
        from src.data.data_sets import get_dataset_years

        # Convert df back to proper format for get_dataset_years (PDF)
        data_dict_for_years = {
            year: results["dataframe"].loc[year].values
            for year in results["dataframe"].index
        }
        start_year, end_year = get_dataset_years(data_dict_for_years)
        story.append(
            Paragraph(
                f"3. {i18n.t_print('monthly_climatology')} ({start_year}-{end_year})",
                heading_style,
            )
        )

        # Create climatology table data
        climatology_data = [
            [
                i18n.t_print("month_header"),
                i18n.t_print("average_header"),
                i18n.t_print("maximum_header"),
                i18n.t_print("minimum_header"),
                i18n.t_print("std_header"),
            ]
        ]

        for i, month in enumerate(months):
            month_data = df[month]

            # Calculate month statistics for PDF
            max_val = month_data.max()
            min_val = month_data.min()
            std_val = month_data.std()

            climatology_data.append(
                [
                    month,
                    # Access climatology data for PDF
                    f"{climatology.iloc[i]:.4f}",
                    f"{max_val:.4f}",
                    f"{min_val:.4f}",
                    f"{std_val:.4f}",
                ]
            )

        # Create and style table
        climatology_table = Table(
            climatology_data,
            colWidths=[1 * inch, 1 * inch, 1 * inch, 1 * inch, 1 * inch],
        )
        climatology_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                    ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ]
            )
        )

        story.append(climatology_table)
        story.append(Spacer(1, 12))

        # 4. ANOMALY ANALYSIS
        story.append(Paragraph(f"4. {i18n.t_print('anomaly_analysis')}", heading_style))

        # Get anomaly statistics for PDF
        anomaly_stats_data = {
            "mean": annual_anomalies.mean(),
            "std": annual_anomalies.std(),
            "min": annual_anomalies.min(),
            "max": annual_anomalies.max(),
            "idxmax": annual_anomalies.idxmax(),
            "idxmin": annual_anomalies.idxmin(),
        }

        # Count positive/negative anomalies for PDF
        positive_count = sum(annual_anomalies > 0)
        negative_count = sum(annual_anomalies < 0)

        anomaly_stats_text = f"""
        • {i18n.t_print('year_highest_positive_anomaly')}: {anomaly_stats_data['idxmax']} ({anomaly_stats_data['max']:.4f})<br/>
        • {i18n.t_print('year_highest_negative_anomaly')}: {anomaly_stats_data['idxmin']} ({anomaly_stats_data['min']:.4f})<br/>
        • {i18n.t_print('years_positive_anomaly')}: {positive_count}<br/>
        • {i18n.t_print('years_negative_anomaly')}: {negative_count}<br/>
        • {i18n.t_print('standard_deviation_anomalies')}: {anomaly_stats_data['std']:.4f}
        """
        story.append(Paragraph(anomaly_stats_text, styles["Normal"]))
        story.append(Spacer(1, 12))

        # 5. SEASONALITY
        story.append(Paragraph(f"5. {i18n.t_print('seasonality')}", heading_style))

        import numpy as np

        # Process climatology for PDF
        clim_stats = {
            "mean": climatology.mean(),
            "std": climatology.std(),
            "min": climatology.min(),
            "max": climatology.max(),
            "idxmax": climatology.idxmax(),
            "idxmin": climatology.idxmin(),
        }
        seasonal_amplitude = clim_stats["max"] - clim_stats["min"]

        # Find max and min months for PDF
        max_month = months[np.argmax(climatology.values)]
        min_month = months[np.argmin(climatology.values)]

        seasonality_stats_text = f"""
        • {i18n.t_print('month_highest_irradiance')}: {max_month} ({clim_stats['max']:.4f} {i18n.t_print('kwh_per_m2_day')})<br/>
        • {i18n.t_print('month_lowest_irradiance')}: {min_month} ({clim_stats['min']:.4f} {i18n.t_print('kwh_per_m2_day')})<br/>
        • {i18n.t_print('seasonal_amplitude')}: {seasonal_amplitude:.4f} {i18n.t_print('kwh_per_m2_day')}
        """
        story.append(Paragraph(seasonality_stats_text, styles["Normal"]))
        story.append(Spacer(1, 12))

        # 6. INTERANNUAL VARIABILITY
        story.append(
            Paragraph(f"6. {i18n.t_print('interannual_variability')}", heading_style)
        )

        # Use the already calculated annual_stats_data
        cv = (annual_stats_data["std"] / annual_stats_data["mean"]) * 100
        interannual_range = annual_stats_data["max"] - annual_stats_data["min"]

        variability_stats_text = f"""
        • {i18n.t_print('annual_coefficient_variation')}: {cv:.2f}%<br/>
        • {i18n.t_print('interannual_range')}: {interannual_range:.4f} {i18n.t_print('kwh_per_m2_day')}
        """
        story.append(Paragraph(variability_stats_text, styles["Normal"]))
        story.append(Spacer(1, 12))

        # 7. FINAL SUMMARY
        story.append(Paragraph(f"7. {i18n.t_print('final_summary')}", heading_style))

        trend_direction = (
            i18n.t_print("negative")
            if trend_coefficients[0] < 0
            else i18n.t_print("positive")
        )

        summary_stats_text = f"""
        • {i18n.t_print('general_trend')}: {trend_direction.upper()}<br/>
        • {i18n.t_print('seasonal_variability')}: {seasonal_amplitude:.4f} {i18n.t_print('kwh_per_m2_day')}<br/>
        • {i18n.t_print('interannual_stability')}: {cv:.2f}% {i18n.t_print('variation')}
        """
        story.append(Paragraph(summary_stats_text, styles["Normal"]))
        story.append(Spacer(1, 12))

    def export_to_docx(self, all_results, dataset_names):
        """
        Export analysis results to DOCX document

        Args:
            all_results (dict): Dictionary with analysis results for all datasets
            dataset_names (list): List of dataset names

        Returns:
            Path: path to the created DOCX file
        """
        # Create document
        doc = Document()

        # Add title
        title = doc.add_heading("", 0)
        title_run = title.runs[0] if title.runs else title.add_run()
        title_run.text = f"{i18n.t_print('solar_irradiance_analysis')}"
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add subtitle with period
        start_year, end_year = get_all_datasets_period()
        subtitle = doc.add_heading(
            f"{i18n.t_print('period')}: {start_year}-{end_year}", level=2
        )
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add generation date
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(
            f"{i18n.t_print('analysis')} {i18n.t_print('period').lower()}: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        date_run.font.size = Pt(10)
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_page_break()

        # Add executive summary
        doc.add_heading(f"{i18n.t_print('comparative_summary')}", level=1)

        # Create summary table
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = i18n.t_print("analysis")
        header_cells[1].text = (
            f"{i18n.t_print('mean_irradiance')} ({i18n.t_print('kwh_per_m2_day')})"
        )
        header_cells[2].text = (
            f"{i18n.t_print('trend')} ({i18n.t_print('kwh_per_m2_year')})"
        )
        header_cells[3].text = f"{i18n.t_print('variability')} (%)"

        # Data rows
        for name, results in all_results.items():
            row_cells = table.add_row().cells
            trend = results["trend_coefficients"][0]
            annual_stats = {
                "mean": results["annual_mean"].mean(),
                "std": results["annual_mean"].std(),
                "min": results["annual_mean"].min(),
                "max": results["annual_mean"].max(),
                "idxmax": results["annual_mean"].idxmax(),
                "idxmin": results["annual_mean"].idxmin(),
            }
            mean_irradiance = annual_stats["mean"]
            cv = (annual_stats["std"] / annual_stats["mean"]) * 100

            row_cells[0].text = name
            row_cells[1].text = f"{mean_irradiance:.4f}"
            row_cells[2].text = f"{trend:.6f}"
            row_cells[3].text = f"{cv:.2f}%"

        # Add comparative plot if multiple datasets
        if len(all_results) > 1:
            doc.add_paragraph()
            comparative_plot_path = self.create_comparative_plot(all_results)
            if comparative_plot_path:
                doc.add_picture(str(comparative_plot_path), width=Inches(7))
                doc.add_paragraph()

        # Add detailed analysis for each dataset
        for name, results in all_results.items():
            doc.add_page_break()
            doc.add_heading(f"{i18n.t_print('detailed_statistics')} - {name}", level=1)

            # Get location from dataset metadata if available
            months = i18n.get_months_print()
            from src.data.data_sets import DATASET_METADATA, get_dataset_display_name

            location = "Location data"  # Default fallback

            # Try to match by both original name and translated name
            for dataset_key, metadata in DATASET_METADATA.items():
                original_name = metadata["name"]
                translated_name = get_dataset_display_name(dataset_key, i18n)
                if original_name == name or translated_name == name:
                    location = metadata["location"]
                    break

            # Add location as subtitle
            if location != "Location data":
                if is_combined_dataset(name, location):
                    formatted_location = format_location_for_display(
                        location, True, i18n
                    )
                    # For document, format each station on a new line
                    location_para = doc.add_paragraph()
                    location_run = location_para.add_run(formatted_location)
                    location_run.font.size = Pt(12)
                    location_run.italic = True
                    location_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    location_para = doc.add_paragraph()
                    location_run = location_para.add_run(location)
                    location_run.font.size = Pt(12)
                    location_run.italic = True
                    location_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            plot_paths = self.create_plots_for_dataset(results, name, months, location)

            # Add main analysis plot to document
            if "main_plot" in plot_paths:
                doc.add_paragraph()
                doc.add_picture(str(plot_paths["main_plot"]), width=Inches(6))
                doc.add_paragraph()

            # Add decade analysis plot to document
            if "decade_plot" in plot_paths:
                doc.add_picture(str(plot_paths["decade_plot"]), width=Inches(6))
                doc.add_paragraph()

            # Add complete detailed analysis as shown in console
            self.add_complete_analysis_to_docx(doc, results, name)

        # Save document
        filename = f"solar_irradiance_analysis_{self.timestamp}.docx"
        filepath = self.output_dir / filename
        doc.save(filepath)

        return filepath

    def export_to_pdf(self, all_results, dataset_names):
        """
        Export analysis results to PDF document

        Args:
            all_results (dict): Dictionary with analysis results for all datasets
            dataset_names (list): List of dataset names

        Returns:
            Path: path to the created PDF file
        """
        filename = f"solar_irradiance_analysis_{self.timestamp}.pdf"
        filepath = self.output_dir / filename

        # Create PDF document
        doc = SimpleDocTemplate(str(filepath), pagesize=A4)
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # Center
        )

        heading_style = ParagraphStyle(
            "CustomHeading",
            parent=styles["Heading2"],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.darkblue,
        )

        # Build document content
        story = []

        # Title
        start_year, end_year = get_all_datasets_period()
        story.append(
            Paragraph(f"{i18n.t_print('solar_irradiance_analysis')}", title_style)
        )
        story.append(
            Paragraph(
                f"{i18n.t_print('period')}: {start_year}-{end_year}", styles["Heading2"]
            )
        )
        story.append(Spacer(1, 12))

        # Generation date
        story.append(
            Paragraph(
                f"{i18n.t_print('analysis')} {i18n.t_print('period').lower()}: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                styles["Normal"],
            )
        )
        story.append(Spacer(1, 20))

        # Executive summary
        story.append(Paragraph(f"{i18n.t_print('comparative_summary')}", heading_style))

        # Summary table data
        table_data = [
            [
                i18n.t_print("analysis"),
                f"{i18n.t_print('mean_irradiance')}\n({i18n.t_print('kwh_per_m2_day')})",
                f"{i18n.t_print('trend')}\n({i18n.t_print('kwh_per_m2_year')})",
                f"{i18n.t_print('variability')}\n(%)",
            ]
        ]

        for name, results in all_results.items():
            trend = results["trend_coefficients"][0]
            annual_stats = {
                "mean": results["annual_mean"].mean(),
                "std": results["annual_mean"].std(),
                "min": results["annual_mean"].min(),
                "max": results["annual_mean"].max(),
                "idxmax": results["annual_mean"].idxmax(),
                "idxmin": results["annual_mean"].idxmin(),
            }
            mean_irradiance = annual_stats["mean"]
            cv = (annual_stats["std"] / annual_stats["mean"]) * 100

            table_data.append(
                [name, f"{mean_irradiance:.4f}", f"{trend:.6f}", f"{cv:.2f}%"]
            )

        # Create table
        table = Table(
            table_data, colWidths=[2 * inch, 1.5 * inch, 1.5 * inch, 1 * inch]
        )
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, 0), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                    ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ]
            )
        )

        story.append(table)

        # Add comparative plot if multiple datasets
        if len(all_results) > 1:
            story.append(Spacer(1, 20))
            comparative_plot_path = self.create_comparative_plot(all_results)
            if comparative_plot_path:
                story.append(
                    Image(
                        str(comparative_plot_path), width=7 * inch, height=5.25 * inch
                    )
                )
                story.append(Spacer(1, 20))

        story.append(PageBreak())

        # Detailed analysis for each dataset
        for name, results in all_results.items():
            story.append(
                Paragraph(
                    f"{i18n.t_print('detailed_statistics')} - {name}", title_style
                )
            )

            # Get location from dataset metadata if available
            months = i18n.get_months_print()
            from src.data.data_sets import DATASET_METADATA, get_dataset_display_name

            location = "Location data"  # Default fallback

            # Try to match by both original name and translated name
            for dataset_key, metadata in DATASET_METADATA.items():
                original_name = metadata["name"]
                translated_name = get_dataset_display_name(dataset_key, i18n)
                if original_name == name or translated_name == name:
                    location = metadata["location"]
                    break

            # Add location as subtitle
            if location != "Location data":
                # Create italic style for location
                location_style = ParagraphStyle(
                    "LocationStyle",
                    parent=styles["Normal"],
                    fontSize=12,
                    fontName="Helvetica-Oblique",
                    alignment=1,  # Center alignment
                    spaceAfter=12,
                )

                if is_combined_dataset(name, location):
                    formatted_location = format_location_for_display(
                        location, True, i18n
                    )
                    # Replace newlines with <br/> for PDF
                    formatted_location_pdf = formatted_location.replace("\n", "<br/>")
                    story.append(Paragraph(formatted_location_pdf, location_style))
                else:
                    story.append(Paragraph(location, location_style))
            plot_paths = self.create_plots_for_dataset(results, name, months, location)

            # Add main analysis plot to document
            if "main_plot" in plot_paths:
                story.append(Spacer(1, 12))
                story.append(
                    Image(
                        str(plot_paths["main_plot"]), width=6 * inch, height=4.5 * inch
                    )
                )
                story.append(Spacer(1, 12))

            # Add decade analysis plot to document
            if "decade_plot" in plot_paths:
                story.append(
                    Image(
                        str(plot_paths["decade_plot"]), width=6 * inch, height=3 * inch
                    )
                )
                story.append(Spacer(1, 12))

            # Add complete detailed analysis as shown in console
            self.add_complete_analysis_to_pdf(
                story, styles, heading_style, results, name
            )

            if (
                name != list(all_results.keys())[-1]
            ):  # Add page break except for last dataset
                story.append(PageBreak())

        # Build PDF
        doc.build(story)

        return filepath

    def export_analysis_results(self, all_results, format="both"):
        """
        Export analysis results to document(s)

        Args:
            all_results (dict): Dictionary with analysis results for all datasets
            format (str): Export format - 'docx', 'pdf', or 'both'

        Returns:
            dict: Dictionary with paths to created files
        """
        dataset_names = list(all_results.keys())
        created_files = {}

        if format in ["docx", "both"]:
            try:
                docx_path = self.export_to_docx(all_results, dataset_names)
                created_files["docx"] = docx_path
                print(f"✅ {i18n.t_print('analysis')} DOCX: {docx_path}")
            except Exception as e:
                print(f"❌ {i18n.t_print('error_exporting_reports')} DOCX: {e}")

        if format in ["pdf", "both"]:
            try:
                pdf_path = self.export_to_pdf(all_results, dataset_names)
                created_files["pdf"] = pdf_path
                print(f"✅ {i18n.t_print('analysis')} PDF: {pdf_path}")
            except Exception as e:
                print(f"❌ {i18n.t_print('error_exporting_reports')} PDF: {e}")

        return created_files


def create_reports(all_results, format="both", output_dir="reports"):
    """
    Convenience function to create reports

    Args:
        all_results (dict): Analysis results
        format (str): Export format - 'docx', 'pdf', or 'both'
        output_dir (str): Output directory

    Returns:
        dict: Paths to created files
    """
    exporter = DocumentExporter(output_dir)
    return exporter.export_analysis_results(all_results, format)
